#!/usr/bin/env python
"""
Script pour générer le manuel utilisateur en PDF.
Utilise markdown2 et weasyprint.
"""

import os
import sys
from pathlib import Path

# Ajouter le chemin du projet
BASE_DIR = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(BASE_DIR))

def generate_pdf():
    """Génère le PDF du manuel utilisateur."""
    try:
        import markdown2
        from weasyprint import HTML, CSS
    except ImportError:
        print("Installation des dépendances requises...")
        os.system("pip install markdown2 weasyprint")
        import markdown2
        from weasyprint import HTML, CSS
    
    # Lire le fichier Markdown
    md_path = BASE_DIR / "docs" / "MANUEL_UTILISATEUR.md"
    with open(md_path, "r", encoding="utf-8") as f:
        md_content = f.read()
    
    # Convertir en HTML
    html_content = markdown2.markdown(
        md_content,
        extras=["tables", "fenced-code-blocks", "header-ids", "toc"]
    )
    
    # Style CSS professionnel
    css_style = """
    @page {
        size: A4;
        margin: 2cm 2.5cm;
        @top-center {
            content: "Manuel Utilisateur NetExpress";
            font-size: 10pt;
            color: #666;
        }
        @bottom-center {
            content: counter(page);
            font-size: 10pt;
        }
    }
    
    body {
        font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif;
        font-size: 11pt;
        line-height: 1.6;
        color: #333;
    }
    
    h1 {
        color: #0e6a4c;
        font-size: 24pt;
        border-bottom: 3px solid #0e6a4c;
        padding-bottom: 10px;
        margin-top: 40px;
        page-break-before: always;
    }
    
    h1:first-of-type {
        page-break-before: avoid;
        text-align: center;
        font-size: 32pt;
        border-bottom: none;
    }
    
    h2 {
        color: #0e6a4c;
        font-size: 16pt;
        margin-top: 30px;
        border-left: 4px solid #0e6a4c;
        padding-left: 15px;
    }
    
    h3 {
        color: #333;
        font-size: 13pt;
        margin-top: 20px;
    }
    
    table {
        width: 100%;
        border-collapse: collapse;
        margin: 15px 0;
        font-size: 10pt;
    }
    
    th {
        background-color: #0e6a4c;
        color: white;
        padding: 10px;
        text-align: left;
        font-weight: bold;
    }
    
    td {
        padding: 8px 10px;
        border-bottom: 1px solid #ddd;
    }
    
    tr:nth-child(even) {
        background-color: #f9f9f9;
    }
    
    code {
        background-color: #f4f4f4;
        padding: 2px 6px;
        border-radius: 3px;
        font-family: Consolas, monospace;
        font-size: 10pt;
    }
    
    strong {
        color: #0e6a4c;
    }
    
    ol, ul {
        margin-left: 20px;
    }
    
    li {
        margin: 5px 0;
    }
    
    hr {
        border: none;
        border-top: 1px solid #ddd;
        margin: 30px 0;
    }
    
    blockquote {
        border-left: 4px solid #0e6a4c;
        padding-left: 15px;
        margin: 15px 0;
        color: #666;
        font-style: italic;
    }
    
    .toc {
        background-color: #f9f9f9;
        padding: 20px;
        border-radius: 8px;
        margin: 20px 0;
    }
    
    .toc a {
        color: #0e6a4c;
        text-decoration: none;
    }
    
    /* Emojis et icones */
    .emoji {
        font-size: 14pt;
    }
    """
    
    # HTML complet
    full_html = f"""
    <!DOCTYPE html>
    <html lang="fr">
    <head>
        <meta charset="UTF-8">
        <title>Manuel Utilisateur NetExpress</title>
    </head>
    <body>
        {html_content}
    </body>
    </html>
    """
    
    # Générer le PDF
    output_path = BASE_DIR / "docs" / "MANUEL_UTILISATEUR.pdf"
    
    HTML(string=full_html).write_pdf(
        output_path,
        stylesheets=[CSS(string=css_style)]
    )
    
    print(f"PDF genere avec succes : {output_path}")
    return output_path


def generate_docx():
    """Génère le document Word du manuel utilisateur."""
    try:
        from docx import Document
        from docx.shared import Pt, Inches, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.style import WD_STYLE_TYPE
    except ImportError:
        print("Installation de python-docx...")
        os.system("pip install python-docx")
        from docx import Document
        from docx.shared import Pt, Inches, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.style import WD_STYLE_TYPE
    
    # Lire le fichier Markdown
    md_path = BASE_DIR / "docs" / "MANUEL_UTILISATEUR.md"
    with open(md_path, "r", encoding="utf-8") as f:
        lines = f.readlines()
    
    # Créer le document
    doc = Document()
    
    # Styles
    style = doc.styles['Heading 1']
    style.font.color.rgb = RGBColor(14, 106, 76)  # Vert NetExpress
    style.font.size = Pt(24)
    
    style2 = doc.styles['Heading 2']
    style2.font.color.rgb = RGBColor(14, 106, 76)
    style2.font.size = Pt(16)
    
    style3 = doc.styles['Heading 3']
    style3.font.size = Pt(13)
    
    # Parser le markdown basique
    current_table = None
    table_headers = []
    
    for line in lines:
        line = line.rstrip()
        
        if line.startswith('# '):
            doc.add_heading(line[2:], level=1)
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=3)
        elif line.startswith('| ') and '|' in line:
            # Tableau
            cells = [c.strip() for c in line.split('|')[1:-1]]
            if cells and not all(c.startswith('-') for c in cells):
                if current_table is None:
                    # Nouveau tableau
                    current_table = doc.add_table(rows=1, cols=len(cells))
                    current_table.style = 'Table Grid'
                    hdr_cells = current_table.rows[0].cells
                    for i, cell in enumerate(cells):
                        hdr_cells[i].text = cell
                    table_headers = cells
                else:
                    # Ajouter une ligne
                    row_cells = current_table.add_row().cells
                    for i, cell in enumerate(cells):
                        if i < len(row_cells):
                            row_cells[i].text = cell
        elif line.startswith('---'):
            current_table = None
            doc.add_paragraph()
        elif line.startswith('1. ') or line.startswith('- '):
            # Liste
            text = line[3:] if line.startswith('1. ') else line[2:]
            p = doc.add_paragraph(text, style='List Bullet' if line.startswith('-') else 'List Number')
        elif line.strip():
            current_table = None
            # Texte normal - gérer le gras
            p = doc.add_paragraph()
            parts = line.split('**')
            for i, part in enumerate(parts):
                run = p.add_run(part)
                if i % 2 == 1:  # Parties impaires = gras
                    run.bold = True
        else:
            current_table = None
    
    # Sauvegarder
    output_path = BASE_DIR / "docs" / "MANUEL_UTILISATEUR.docx"
    doc.save(output_path)
    
    print(f"Document Word genere avec succes : {output_path}")
    return output_path


if __name__ == "__main__":
    print("Generation du manuel utilisateur...")
    print()
    
    try:
        pdf_path = generate_pdf()
        print(f"PDF: {pdf_path}")
    except Exception as e:
        print(f"Erreur PDF: {e}")
    
    print()
    
    try:
        docx_path = generate_docx()
        print(f"DOCX: {docx_path}")
    except Exception as e:
        print(f"Erreur DOCX: {e}")

